from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('MC-RCPSP : Problème d’Ordonnancement de Projet à Ressources Limitées et Critères Multiples', 0)

# Add Sections to the Document
doc.add_heading('1. Définition générale', level=1)
doc.add_paragraph(
    "Le RCPSP (Resource-Constrained Project Scheduling Problem) est un problème classique en recherche opérationnelle où l'on cherche à planifier un ensemble de tâches d’un projet, tout en respectant :\n"
    "- Les contraintes de précédence (certaines tâches doivent être terminées avant que d'autres ne commencent).\n"
    "- Les contraintes de ressources limitées (chaque tâche utilise un certain nombre de ressources disponibles en quantité limitée).\n"
    "L'objectif principal est généralement de minimiser la durée totale du projet (makespan)."
)

doc.add_heading('2. Extension vers le MC-RCPSP', level=1)
doc.add_paragraph(
    "Le MC-RCPSP (Multi-Criteria RCPSP) est une extension du RCPSP classique où plusieurs objectifs sont considérés simultanément :\n"
    "Par exemple :\n"
    "- Minimiser la durée totale (makespan).\n"
    "- Réduire le coût global du projet.\n"
    "- Optimiser l'utilisation des ressources.\n"
    "- Réduire le nombre de dépassements de ressources ou équilibrer leur utilisation.\n"
    "Cela le rend plus proche des projets réels, où plusieurs aspects doivent être optimisés ensemble."
)

doc.add_heading('3. Objectif', level=1)
doc.add_paragraph(
    "L'objectif du MC-RCPSP est de trouver un calendrier d’exécution des tâches qui respecte les contraintes de précédence et de ressources, tout en optimisant un ou plusieurs objectifs selon les priorités du projet. L’objectif peut être :\n"
    "- Mono-objectif (makespan minimal).\n"
    "- Multi-objectifs (makespan + coût + ressources, etc.)."
)

doc.add_heading('4. Méthodes de résolution', level=1)
doc.add_paragraph(
    "Plusieurs approches peuvent être utilisées pour résoudre le MC-RCPSP :\n"
    "- Méthodes exactes : programmation linéaire en nombres entiers (ILP), Branch and Bound.\n"
    "- Méthodes heuristiques : listes de priorité, règles simples.\n"
    "- Métaheuristiques : algorithmes génétiques, colonies de fourmis, recuit simulé, etc.\n"
    "- Approches multi-objectifs : NSGA-II, SPEA2, etc."
)

doc.add_heading('5. Formulation mathématique simplifiée', level=1)
doc.add_paragraph(
    "Soit :\n"
    "- N : nombre de tâches\n"
    "- R : nombre de types de ressources\n"
    "- d_i : durée de la tâche i\n"
    "- r_ik : ressource k requise par la tâche i\n"
    "- b_k : capacité disponible de la ressource k\n"
    "- S_i : date de début de la tâche i\n"
    "- P_i : ensemble des tâches précédentes de i\n\n"
    "Objectif : Min {makespan = max(S_i + d_i)}\n\n"
    "Contraintes :\n"
    "1. Précédence : S_j ≥ S_i + d_i ∀ i ∈ N, ∀ j ∈ P_i\n"
    "2. Ressources : ∑ r_ik ≤ b_k ∀ k ∈ R à tout temps t"
)

doc.add_heading('6. Exemple pratique', level=1)
doc.add_paragraph(
    "Prenons l'exemple suivant avec 3 tâches et des ressources limitées :\n\n"
    "- Tâches :\n"
    "  - A : durée 2, ressource 2, coût 5\n"
    "  - B : durée 3, ressource 3, coût 4 (dépend de A)\n"
    "  - C : durée 1, ressource 1, coût 3 (dépend de A)\n\n"
    "- Ressources totales disponibles : 4 unités\n\n"
    "Planification :\n"
    "1. A commence à t=0.\n"
    "2. B et C peuvent commencer après A.\n"
    "3. B utilise plus de ressources et prend donc plus de temps à être planifiée."
)

doc.add_heading('7. Exemple de code en Python', level=1)
doc.add_paragraph("""
class Task:
    def __init__(self, name, duration, resource, cost, predecessors=[]):
        self.name = name
        self.duration = duration
        self.resource = resource
        self.cost = cost
        self.predecessors = predecessors
        self.start = None
        self.end = None

def schedule(tasks, max_resource):
    time = 0
    scheduled = []
    while len(scheduled) < len(tasks):
        ready = [t for t in tasks if t not in scheduled and all(p in [s.name for s in scheduled] for p in t.predecessors)]
        ready.sort(key=lambda x: x.cost)
        for t in ready:
            if t.resource <= max_resource:
                t.start = time
                t.end = time + t.duration
                scheduled.append(t)
                time = t.end
                break
    makespan = max(t.end for t in scheduled)
    total_cost = sum(t.cost * t.duration for t in scheduled)
    return makespan, total_cost
""")

doc.add_heading('8. Exemple de code en C++', level=1)
doc.add_paragraph("""
#include <iostream>
#include <vector>
#include <algorithm>
#include <string>
using namespace std;

struct Task {
    string name;
    int duration, resource, cost;
    vector<string> predecessors;
    int start = -1, end = -1;
};

bool ready(const Task& task, const vector<string>& done) {
    for (auto& p : task.predecessors)
        if (find(done.begin(), done.end(), p) == done.end()) return false;
    return true;
}

int main() {
    vector<Task> tasks = {
        {"A", 2, 2, 5, {}}, 
        {"B", 3, 3, 4, {"A"}}, 
        {"C", 1, 1, 3, {"A"}}
    };
    int time = 0, max_resource = 4;
    vector<string> done;
    while (done.size() < tasks.size()) {
        for (Task& t : tasks) {
            if (t.start == -1 && ready(t, done) && t.resource <= max_resource) {
                t.start = time;
                t.end = time + t.duration;
                done.push_back(t.name);
                time = t.end;
                break;
            }
        }
    }
    return 0;
}
""")

doc.add_heading('9. Conclusion', level=1)
doc.add_paragraph(
    "Le MC-RCPSP est un problème complexe mais très réaliste pour la planification de projets avec des contraintes réelles. "
    "Les exemples fournis ici sont simplifiés, mais ils reflètent bien la logique de base. Des outils avancés permettent d’aborder des projets à grande échelle avec de multiples objectifs à équilibrer."
)

# Save the document
docx_path = "/home/zeiny/Documents/MC-RCPSP/MC_RCPSP_Complete_Explained.docx"
doc.save(docx_path)

docx_path
